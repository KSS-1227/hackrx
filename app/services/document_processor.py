"""
Document processing service for downloading and extracting text from various document formats
"""

import asyncio
import aiofiles
import httpx
import tempfile
import os
from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger
import time

# Document processing imports
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
import textstat
from bs4 import BeautifulSoup

from app.core.config import settings
from app.core.exceptions import DocumentProcessingError, DocumentDownloadError
from app.models.document import DocumentChunk


class DocumentProcessor:
    """Service for processing documents from URLs"""
    
    def __init__(self):
        self.max_size_bytes = settings.MAX_DOCUMENT_SIZE_MB * 1024 * 1024
        self.supported_formats = settings.supported_formats_list
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    async def process_documents(self, document_urls: List[str]) -> List[DocumentChunk]:
        """
        Process multiple documents concurrently
        
        Args:
            document_urls: List of document URLs to process
            
        Returns:
            List of DocumentChunk objects
        """
        semaphore = asyncio.Semaphore(settings.MAX_CONCURRENT_DOWNLOADS)
        
        async def process_single_document(url: str) -> List[DocumentChunk]:
            async with semaphore:
                return await self._process_single_document(url)
        
        # Process documents concurrently
        tasks = [process_single_document(url) for url in document_urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Flatten results and filter out exceptions
        all_chunks = []
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Document processing failed: {str(result)}")
                continue
            all_chunks.extend(result)
        
        logger.info(f"Successfully processed {len(all_chunks)} chunks from {len(document_urls)} documents")
        return all_chunks
    
    async def _process_single_document(self, url: str) -> List[DocumentChunk]:
        """Process a single document from URL"""
        start_time = time.time()
        
        try:
            # Download document
            logger.info(f"Downloading document from: {url}")
            file_path, metadata = await self._download_document(url)
            
            # Extract text based on file format
            logger.info(f"Extracting text from: {metadata['filename']}")
            text_content = await self._extract_text(file_path, metadata['format'])
            
            # Clean up temporary file
            os.unlink(file_path)
            
            # Chunk the text
            chunks = self._chunk_text(text_content, url, metadata)
            
            processing_time = time.time() - start_time
            metadata['processing_time'] = processing_time
            
            logger.info(f"Processed document {metadata['filename']} in {processing_time:.2f}s, created {len(chunks)} chunks")
            
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process document {url}: {str(e)}")
            raise DocumentProcessingError(f"Failed to process document {url}: {str(e)}")
    
    async def _download_document(self, url: str) -> tuple[str, Dict[str, Any]]:
        """Process document from URL or local file path and return file path and metadata"""
        
        try:
            # Check if this is a local file path
            if os.path.isfile(url):
                # This is a local file, not a URL
                file_path = url
                file_size = os.path.getsize(file_path)
                
                # Check file size
                if file_size > self.max_size_bytes:
                    raise DocumentDownloadError(f"Document too large: {file_size} bytes")
                
                # Determine file format
                filename = os.path.basename(file_path)
                file_format = self._determine_format(filename, "")
                
                if file_format not in self.supported_formats:
                    raise DocumentDownloadError(f"Unsupported file format: {file_format}")
                
                metadata = {
                    'filename': filename,
                    'format': file_format,
                    'size_bytes': file_size,
                    'content_type': f"application/{file_format}",
                    'url': url
                }
                
                return file_path, metadata
                
            # If not a local file, treat as URL
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()
                
                # Check file size
                content_length = response.headers.get('content-length')
                if content_length and int(content_length) > self.max_size_bytes:
                    raise DocumentDownloadError(f"Document too large: {content_length} bytes")
                
                # Determine file format
                content_type = response.headers.get('content-type', '').lower()
                filename = self._extract_filename_from_url(url, content_type)
                file_format = self._determine_format(filename, content_type)
                
                if file_format not in self.supported_formats:
                    raise DocumentDownloadError(f"Unsupported file format: {file_format}")
                
                # Save to temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_format}") as temp_file:
                    temp_file.write(response.content)
                    temp_path = temp_file.name
                
                metadata = {
                    'filename': filename,
                    'format': file_format,
                    'size_bytes': len(response.content),
                    'content_type': content_type,
                    'url': url
                }
                
                return temp_path, metadata
                
        except httpx.HTTPError as e:
            raise DocumentDownloadError(f"Failed to download document: {str(e)}")
        except Exception as e:
            raise DocumentDownloadError(f"Unexpected error downloading document: {str(e)}")
    
    def _extract_filename_from_url(self, url: str, content_type: str) -> str:
        """Extract filename from URL or generate one based on content type"""
        path = Path(url)
        if path.suffix:
            return path.name
        
        # Generate filename based on content type
        if 'pdf' in content_type:
            return f"document_{hash(url) % 10000}.pdf"
        elif 'word' in content_type or 'docx' in content_type:
            return f"document_{hash(url) % 10000}.docx"
        elif 'html' in content_type:
            return f"document_{hash(url) % 10000}.html"
        else:
            return f"document_{hash(url) % 10000}.txt"
    
    def _determine_format(self, filename: str, content_type: str) -> str:
        """Determine file format from filename and content type"""
        extension = Path(filename).suffix.lower().lstrip('.')
        
        if extension in self.supported_formats:
            return extension
        
        # Fallback to content type
        if 'pdf' in content_type:
            return 'pdf'
        elif 'word' in content_type or 'docx' in content_type:
            return 'docx'
        elif 'html' in content_type:
            return 'html'
        else:
            return 'txt'
    
    async def _extract_text(self, file_path: str, file_format: str) -> str:
        """Extract text from document based on format"""
        
        try:
            if file_format == 'pdf':
                return await self._extract_pdf_text(file_path)
            elif file_format in ['docx', 'doc']:
                return await self._extract_docx_text(file_path)
            elif file_format == 'html':
                return await self._extract_html_text(file_path)
            elif file_format == 'txt':
                return await self._extract_txt_text(file_path)
            else:
                raise DocumentProcessingError(f"Unsupported format for text extraction: {file_format}")
                
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract text from {file_format} file: {str(e)}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using multiple methods for best results"""
        text_content = ""
        
        try:
            # Method 1: PyMuPDF (fast and good for most PDFs)
            doc = fitz.open(file_path)
            for page in doc:
                text_content += page.get_text() + "\n"
            doc.close()
            
            # If PyMuPDF didn't extract much text, try pdfplumber
            if len(text_content.strip()) < 100:
                logger.info("PyMuPDF extracted minimal text, trying pdfplumber...")
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract DOCX text: {str(e)}")
    
    async def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = ' '.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract HTML text: {str(e)}")
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                return await file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    return await file.read()
            except Exception as e:
                raise DocumentProcessingError(f"Failed to read text file: {str(e)}")
    
    def _chunk_text(self, text: str, source: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Split text into chunks for processing"""
        
        if not text.strip():
            logger.warning(f"No text content extracted from {source}")
            return []
        
        # Simple sentence-aware chunking
        sentences = text.split('. ')
        chunks = []
        current_chunk = ""
        chunk_id = 0
        char_position = 0
        
        for sentence in sentences:
            # Add sentence to current chunk
            test_chunk = current_chunk + sentence + ". "
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                # Current chunk is full, save it and start new one
                if current_chunk.strip():
                    start_char = char_position - len(current_chunk)
                    end_char = char_position
                    
                    chunks.append(DocumentChunk(
                        id=f"{hash(source)}_{chunk_id}",
                        content=current_chunk.strip(),
                        source=source,
                        chunk_index=chunk_id,
                        start_char=start_char,
                        end_char=end_char,
                        metadata={
                            **metadata,
                            'chunk_id': chunk_id,
                            'chunk_size': len(current_chunk),
                            'readability_score': textstat.flesch_reading_ease(current_chunk)
                        }
                    ))
                    chunk_id += 1
                
                # Start new chunk with overlap
                if len(chunks) > 0 and self.chunk_overlap > 0:
                    # Take last few sentences for overlap
                    overlap_text = '. '.join(current_chunk.split('. ')[-2:])
                    current_chunk = overlap_text + ". " + sentence + ". "
                else:
                    current_chunk = sentence + ". "
            
            # Update character position
            char_position += len(sentence) + 2  # +2 for the '. ' separator
        
        # Add final chunk
        if current_chunk.strip():
            start_char = char_position - len(current_chunk)
            end_char = char_position
            
            chunks.append(DocumentChunk(
                id=f"{hash(source)}_{chunk_id}",
                content=current_chunk.strip(),
                source=source,
                chunk_index=chunk_id,
                start_char=start_char,
                end_char=end_char,
                metadata={
                    **metadata,
                    'chunk_id': chunk_id,
                    'chunk_size': len(current_chunk),
                    'readability_score': textstat.flesch_reading_ease(current_chunk)
                }
            ))
        
        logger.info(f"Created {len(chunks)} chunks from document {metadata.get('filename', source)}")
        return chunks